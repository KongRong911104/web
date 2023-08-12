from docx.shared import Cm
import docx
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
import time as t
import sys
import os
import json as js
t0=t.time()
doc = docx.Document('./各類作業及考試繳交紀錄表.docx')
table = doc.tables[0]  # 第0个表格

def font_style(point, front, back, chinese=0, underline=0):
    point.text = front
    paragraph = point.paragraphs[0]
    run = paragraph.runs[0]
    run_element = run._element
    rpr = run_element.get_or_add_rPr()
    rFonts = rpr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), "標楷體")
    run.underline = False
    run.font.size = Pt(12)
    new_run = paragraph.add_run(back)
    if chinese == 1:
        run_element = new_run._element
        rpr = run_element.get_or_add_rPr()
        rFonts = rpr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), "標楷體")
    else:
        new_run.font.name = 'Times New Roman'
    new_run.font.size = Pt(12)
    if underline == 1:
        new_run.underline = True
def set_row_height(table, row_index, height_cm):
    row = table.rows[row_index]  # 行的索引，根据实际情况进行调整
    row.height = Cm(height_cm)  # 设置行高，单位为厘米
def set_row_width(table, row_index,i, width_cm):
    row = table.cell(row_index,i)  # 行的索引，根据实际情况进行调整
    row.width = Cm(width_cm) # 设置行高，单位为厘米
def add_rows(table, num_rows):
    for i in range(num_rows):
        table.add_row()
        set_row_height(table, 5+i, 0.8)

#輸入區
file_name = sys.argv[2]
ter=file_name.split("_")
# print(ter)
year=ter[0]
semester=""
if ter[1]=="A":
    semester="上學期"
elif ter[1]=="B":
    semester="下學期"
elif ter[1]=="C":
    semester="暑修上"
elif ter[1]=="D":
    semester="暑修下"
curr_code=ter[2]
data_types=["","小考","期中考","期末考","作業","書面報告","口頭報告","實作成品","口試","其他"]
data_type=data_types[int(ter[3])]
times=ter[4].split(".")[0]
name = sys.argv[3]
ta = sys.argv[4]
day = str(sys.argv[5])
docx_file_name = file_name.split(".")[0]+".docx"
insent_text="55"
row_add=0

# #爬蟲區
cs = {}
input_file="../csReview/spider/"+str(year)+"_"+"1.txt"
with open(input_file, "r",encoding="utf-8") as file:
    for line in file:
        key, values = line.strip().split(":")
        value_list = values.split(",")
        cs[int(key)] = value_list
input_file="../csReview/spider/"+str(year)+"_"+"2.txt"
with open(input_file, "r",encoding="utf-8") as file:
    for line in file:
        key, values = line.strip().split(":")
        value_list = values.split(",")
        cs[int(key)] = value_list
sem = cs[int(curr_code)][0]
cour_name =cs[int(curr_code)][2]

import random
#pandas區
df = pd.read_excel(sys.argv[1], usecols=[0,1,2], engine='openpyxl', dtype={
                               "學號": 'str', "姓名": 'str',"成績": 'int32'})
os.remove(sys.argv[1])
# 按照 "成績" 列进行降序排序
df_sorted = df.sort_values(by="成績", ascending=False)
total_mean = df["成績"].mean()

# 按照 "成績" 列进行降序排序
df_sorted = df.sort_values(by="成績", ascending=False)

# 使用 groupby 将 DataFrame 拆分为三个新的 DataFrame
grouped = df_sorted.groupby(pd.cut(df_sorted["成績"], bins=3, labels=False, duplicates='drop'))

# 创建一个空列表来存储结果
result_list = []

# 处理每个拆分后的 DataFrame

# 处理每个拆分后的 DataFrame
for _, group in grouped:
    # 获取每份中分数最高的两条数据的所有值
    top_two_values = group.nlargest(2, "成績").values.tolist()

    # 获取该份的长度、平均值、最高分和最低分
    length = len(group)
    mean = group["成績"].mean()
    max_score = group["成績"].max()
    min_score = group["成績"].min()

    # 将结果存储到列表中
    result_list.append((top_two_values, length, mean, max_score, min_score))
docx_result = []
# 打印结果列表
for i, result in enumerate(result_list):
    top_two_values, length, mean, max_score, min_score = result
    docx_result.append("訖： "+str(max_score))
    docx_result.extend(top_two_values[0])
    docx_result.append("起： "+str(min_score))
    docx_result.extend(top_two_values[1])
    docx_result.append(length)
    docx_result.append(mean)
docx_result.append(round(total_mean,1))

shi_text=str(len(df))
cha_text=str(int(insent_text)-len(df))

#前處理區
s = {"必修": "\u2610必修　", "必選": "\u2610必選　", "選修": "\u2610選修"}
s[sem] = "\u2611"+s[sem][1:]
type__ = "　　 　　　"
type_ = {"小考": "\u2610 小考　", "期中考": "\u2610 期中考　", "期末考": "\u2610 期末考　　", "作業": "\u2610 作業　　", "書面報告": "\u2610 書面報告\n", "口頭報告": "\u2610 口頭報告　", "實作成品":  "\u2610 實作成品　", "口試": "\u2610 口試　", "其他": "\u2610 其他："}
type_[data_type]="\u2611"+type_[data_type][1:]
type_text=type_["小考"]+type_["期中考"]+type_["期末考"]+type_["作業"]+type_["書面報告"]+type_["口頭報告"]+type_["實作成品"]+type_["口試"]+type_["其他"]
space_num = 10-len(cour_name)
if space_num < 0:
    space_num = 0
cour_name_space = "　"*space_num
times_font="  　　第"
times_="　"+times+"　"
times_back="次"

#處理區
font_style(table.cell(0, 0), "學　　年：", "　　"+year+"　　", underline=1)
font_style(table.cell(0, 2), "學　　期：", "　"+semester+"　　", 1, 1)
font_style(table.cell(0, 4), " 課程類別 ：", s["必修"]+s["必選"]+s["選修"], 1)
set_row_height(table, 0, 1.19)
font_style(table.cell(1, 0), "授課教師：", "  "+name+" 　", 1, 1)
font_style(table.cell(1, 2), "科目代號：", "　"+curr_code+"　　", underline=1)
font_style(table.cell(1, 4), "科目名稱：", cour_name_space +
           cour_name+cour_name_space, 1, 1)
set_row_height(table, 1, 1.19)
font_style(table.cell(2, 0), "助　　教：", "  "+ta+" 　", 1, 1)
font_style(table.cell(2, 2), "評量日期：", " "+day+"　", underline=1)
font_style(table.cell(2, 4), "電子檔名稱：", "　"+docx_file_name+"　　", underline=1)
set_row_height(table, 2, 1.19)

table.cell(3, 4).text=type_text
paragraph = table.cell(3, 4).paragraphs[0]
run = paragraph.runs[0]
run_element = run._element
rpr = run_element.get_or_add_rPr()
rFonts = rpr.get_or_add_rFonts()
rFonts.set(qn('w:eastAsia'), "標楷體")
run.underline = False
run.font.size = Pt(12)
new_run = paragraph.add_run(type__)
new_run.font.size = Pt(12)
new_run.underline = True
new_run = paragraph.add_run(times_font)
new_run.font.size = Pt(12)
run_element = new_run._element
rpr = run_element.get_or_add_rPr()
rFonts = rpr.get_or_add_rFonts()
rFonts.set(qn('w:eastAsia'), "標楷體")
new_run = paragraph.add_run(times_)
new_run.font.size = Pt(12)
new_run.underline = True
new_run = paragraph.add_run(times_back)
new_run.font.size = Pt(12)
run_element = new_run._element
rpr = run_element.get_or_add_rPr()
rFonts = rpr.get_or_add_rFonts()
rFonts.set(qn('w:eastAsia'), "標楷體")
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
table.cell(3,4).vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
set_row_height(table, 3, 2.06)

insent_front="應交(考)份數："
insent="  "+insent_text+"  "
insent_back=" 份"

table.cell(4, 0).text=insent_front
paragraph = table.cell(4, 0).paragraphs[0]
run = paragraph.runs[0]
run_element = run._element
rpr = run_element.get_or_add_rPr()
rFonts = rpr.get_or_add_rFonts()
rFonts.set(qn('w:eastAsia'), "標楷體")
run.underline = False
run.font.size = Pt(12)
new_run = paragraph.add_run(insent)
new_run.font.size = Pt(12)
new_run.underline = True
new_run = paragraph.add_run(insent_back)
new_run.font.size = Pt(12)
run_element = new_run._element
rpr = run_element.get_or_add_rPr()
rFonts = rpr.get_or_add_rFonts()
rFonts.set(qn('w:eastAsia'), "標楷體")

insent_front="實交(考)份數："
insent="  "+shi_text+"  "
insent_back=" 份"

table.cell(4, 3).text=insent_front
paragraph = table.cell(4, 3).paragraphs[0]
run = paragraph.runs[0]
run_element = run._element
rpr = run_element.get_or_add_rPr()
rFonts = rpr.get_or_add_rFonts()
rFonts.set(qn('w:eastAsia'), "標楷體")
run.underline = False
run.font.size = Pt(12)
new_run = paragraph.add_run(insent)
new_run.font.size = Pt(12)
new_run.underline = True
new_run = paragraph.add_run(insent_back)
new_run.font.size = Pt(12)
run_element = new_run._element
rpr = run_element.get_or_add_rPr()
rFonts = rpr.get_or_add_rFonts()
rFonts.set(qn('w:eastAsia'), "標楷體")

insent_front="缺交(考)份數："
insent="  "+cha_text+"  "
insent_back=" 份"

table.cell(4, 5).text=insent_front
paragraph = table.cell(4, 5).paragraphs[0]
run = paragraph.runs[0]
run_element = run._element
rpr = run_element.get_or_add_rPr()
rFonts = rpr.get_or_add_rFonts()
rFonts.set(qn('w:eastAsia'), "標楷體")
run.underline = False
run.font.size = Pt(12)
new_run = paragraph.add_run(insent)
new_run.font.size = Pt(12)
new_run.underline = True
new_run = paragraph.add_run(insent_back)
new_run.font.size = Pt(12)
run_element = new_run._element
rpr = run_element.get_or_add_rPr()
rFonts = rpr.get_or_add_rFonts()
rFonts.set(qn('w:eastAsia'), "標楷體")
set_row_height(table, 4, 0.79)

table = doc.tables[1]

add_rows(table, row_add)

table = doc.tables[2]
index=0
for i in range(1, int(0.2*len(docx_result)+0.8)):
    for j in range(2, 6):
        if j==3:
            if i%2:
                tt="2."
            else:
                tt="1."
            table.cell(i,j).text=tt+str(docx_result[index])
        else:
            
            table.cell(i,j).text=str(docx_result[index])
            if j>3:
                table.cell(i,j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        index+=1
    if (i%2==0):
        for j in range(6, 8):  
            
            table.cell(i-1,j).text=str(docx_result[index])
            table.cell(i-1,j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            index+=1
table.cell(7,7).text=str(docx_result[-1])
table.cell(7,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
# 獲取最後一個頁的頁腳
footer = doc.sections[0].footer

table=footer.tables[0]
# print(table.cell(0, 6).text)
set_row_height(table,0,0.74)
set_row_width(table,0,0,2.6)
set_row_width(table,0,3,2.6)
set_row_width(table,0,6,2.6)
set_row_width(table,0,1,3.86)
set_row_width(table,0,4,3.86)
set_row_width(table,0,7,2.05)

doc.save("./word_output/{}/{}.docx".format(sys.argv[6],sys.argv[2]))
with open('xxx.txt', 'w') as f:
    f.write('{}\n'.format(str(t.time()-t0)))