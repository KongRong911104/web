import os
import sys
import json as js
# pip install fake-useragent
from fake_useragent import UserAgent
# pip install pandas
# sudo pip3 install openpyxl
import pandas as pd
# import openpyxl
# pip install python-docx
import docx
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
# pip install urllib3
import urllib.request as req
# pip install beautifulsoup4
import bs4
# pip install fake-useragent
from fake_useragent import UserAgent
ua = UserAgent()

year = str(sys.argv[2])
semester = ""
if sys.argv[3] == "上":
    semester = "1"
elif sys.argv[3] == "下":
    semester = "2"
else:
    print(js.dumps("錯誤"))


def spider(PATH, part=0, cs={}):

    request = req.Request(PATH, headers={"User-Agent": ua.random})
    with req.urlopen(request) as response:
        data = response.read().decode("utf-8")
    root = bs4.BeautifulSoup(data, "html.parser")
    titles = root.find_all("td", {"data-title": "課程名稱"})
    y = 0
    c = ""
    for i in titles:
        sn = int(str(i).replace("\n", '').strip().replace(
            '<td data-title="課程名稱"><a href="/view/{}'.format(str(year)+"/"+str(semester)+"/"), '').split('"')[0])
        text = i.text.strip().split('-')[0]
        if (part == 0):
            if (y == 0):
                c = str(sn)[0]
                y = 1
            if (str(sn)[0] == c):
                cs[sn] = ([text, "日間部"])
            else:
                cs[sn] = ([text, "碩士班"])
        else:
            cs[sn] = ([text, "碩士在職專班"])
    return cs


cs = spider(
    "https://course.thu.edu.tw/view-dept/{}/{}/350".format(year, semester))
cs = spider(
    "https://course.thu.edu.tw/view-dept/{}/{}/357".format(year, semester), 1, cs)


def font_style(a, string, b=0):
    a.paragraphs[0].style.font.name = string
    if b == 1:
        a.paragraphs[0].style._element.rPr.rFonts.set(qn('w:eastAsia'), string)


def font_color(a, b=0):
    for run in a.paragraphs[0].runs:
        if b == 2:
            run.font.color.rgb = RGBColor(255, 0, 0)
        else:
            run.font.color.rgb = RGBColor(0, 0, 255)
        if b == 1:
            run.font.size = Pt(13)
        else:
            run.font.size = Pt(18)


file = sys.argv[1].split("::123::")[:-1]

for i in range(len(file)):
    if file[i].split('.')[-1] != 'xlsx':
        print(js.dumps("副檔名格式錯誤 {}".format(file[i].split('.')[-1])))
if semester == "1":
    semester = "上"
elif semester == "2":
    semester = "下"
else:
    print(js.dumps("學期格式錯誤"))
output = {}
readExcelFile = pd.read_excel(file[0], usecols=[7, 11, 10, 15, 17], engine='openpyxl', dtype={
                              "score": 'int32', "curr_code": 'int32', "cour_cname": 'str', "curr_cname": 'str', 'teac_name': 'str'})

for i in range(1, len(file)):
    readExcelFile2 = pd.read_excel(file[i], usecols=[7, 11, 10, 15, 17], engine='openpyxl', dtype={
        "score": 'int32', "curr_code": 'int32', "cour_cname": 'str', "curr_cname": 'str', 'teac_name': 'str'})
    os.remove(file[i])
    readExcelFile = pd.concat([readExcelFile, readExcelFile2], axis=0)
ExcelFile = pd.DataFrame(readExcelFile)

teac_name_set = readExcelFile["teac_name"].unique()
curr_code_set = readExcelFile["curr_code"].unique()
data = []
for i in curr_code_set:
    data.append(ExcelFile.loc[ExcelFile['curr_code'] == i])

a = []
for i in range(len(teac_name_set)):
    for j in range(len(data)):
        if data[j]["teac_name"].iloc[0] == teac_name_set[i]:
            a.append(data[j])
folder_path = '../word_output'
index = 0

for i in range(len(a)):

    curr_code = a[i].iloc[0, 2]
    compulsory = cs[curr_code][0]
    academic = cs[curr_code][1]
    cour_cname = a[i].iloc[0, 1]
    curr_cname = a[i].iloc[0, 3]
    teac_name = a[i].iloc[0, 4].split(' ')[-1][0:3]
    avg_score = a[i].loc[a[i]['score'] > 20]
    avg_score = round(avg_score['score'].mean(), 1)
    num = len(a[i])
    under_60 = a[i].loc[a[i]['score'] < 60]
    under_20 = under_60.loc[under_60['score'] <= 20]
    # print(curr_code,cour_cname,teac_name,num,avg_score,len(under_60),len(under_20))
    doc = docx.Document('../教學成效表_空.docx')
    tables = doc.tables
    # 老師
    tables[0].cell(0, 1).text = teac_name
    font_color(tables[0].cell(0, 1), 1)
    font_style(tables[0].cell(0, 1), '標楷體', 1)
    # 學年
    tables[0].cell(1, 1).text = year
    tables[0].cell(1, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    font_style(tables[0].cell(1, 1), 'Times New Roman')
    font_color(tables[0].cell(1, 1))
    # 學期
    tables[0].cell(1, 4).text = semester
    font_color(tables[0].cell(1, 4))
    font_style(tables[0].cell(1, 4), '標楷體', 1)
    tables[0].cell(1, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 課程代號
    tables[0].cell(1, 6).text = str(curr_code)
    font_color(tables[0].cell(1, 6))
    font_style(tables[0].cell(1, 6), 'Times New Roman')
    tables[0].cell(1, 6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 課程名稱
    tables[0].cell(2, 1).text = cour_cname
    font_color(tables[0].cell(2, 1), 1)
    font_style(tables[0].cell(2, 1), '標楷體', 1)
    # 學制
    tables[0].cell(2, 6).text = academic
    font_color(tables[0].cell(2, 6), 1)
    font_style(tables[0].cell(2, 6), '標楷體', 1)
    tables[0].cell(2, 6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 修習年級
    tables[0].cell(3, 1).text = curr_cname
    font_color(tables[0].cell(3, 1), 1)
    font_style(tables[0].cell(3, 1), '標楷體', 1)
    tables[0].cell(3, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 學期
    tables[0].cell(3, 4).text = compulsory
    font_color(tables[0].cell(3, 4), 1)
    font_style(tables[0].cell(3, 4), '標楷體', 1)
    tables[0].cell(3, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 平均
    tables[0].cell(3, 6).text = str(avg_score)
    font_color(tables[0].cell(3, 6))
    font_style(tables[0].cell(3, 6), 'Times New Roman')
    tables[0].cell(3, 6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 人數
    tables[0].cell(4, 1).text = str(num)
    font_color(tables[0].cell(4, 1))
    font_style(tables[0].cell(4, 1), 'Times New Roman')
    tables[0].cell(4, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 60以下
    tables[0].cell(4, 4).text = str(len(under_60))
    font_color(tables[0].cell(4, 4))
    font_style(tables[0].cell(4, 4), 'Times New Roman')
    tables[0].cell(4, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 20以下
    tables[0].cell(4, 6).text = str(len(under_20))
    font_color(tables[0].cell(4, 6))
    font_style(tables[0].cell(4, 6), 'Times New Roman')
    tables[0].cell(4, 6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tables[0].cell(5, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    if (len(under_60)/num >= 0.4):
        tables[0].cell(5, 1).text = "✓"
        font_color(tables[0].cell(5, 1), 2)
        font_style(tables[0].cell(5, 1), 'Times New Roman')
        tables[0].cell(
            5, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if avg_score < 60:
        tables[0].cell(6, 1).text = "✓"
        font_color(tables[0].cell(6, 1), 2)
        font_style(tables[0].cell(6, 1), 'Times New Roman')
        tables[0].cell(
            6, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    output[i] = str(curr_code)+"_"+tables[0].cell(2, 1).text + \
        "_"+tables[0].cell(0, 1).text
    if i != 0:
        doc.save(r"../word_output/{}_教學成效檢討回饋表.docx".format(
            str(curr_code)+"_"+tables[0].cell(2, 1).text+"_"+tables[0].cell(0, 1).text))
        new_doc = docx.Document(
            "../word_output/{}教學成效表.docx".format(year+sys.argv[3]))
        dd = docx.Document(r"../word_output/{}_教學成效檢討回饋表.docx".format(
            str(curr_code)+"_"+tables[0].cell(2, 1).text+"_"+tables[0].cell(0, 1).text))
        for element in dd.element.body:
            new_doc.element.body.append(element)
        new_doc.save("../word_output/{}教學成效表.docx".format(year+sys.argv[3]))
        os.remove(r"../word_output/{}_教學成效檢討回饋表.docx".format(
            str(curr_code)+"_"+tables[0].cell(2, 1).text+"_"+tables[0].cell(0, 1).text))

    if i == 0:
        doc.save("../word_output/{}教學成效表.docx".format(year+sys.argv[3]))
print(js.dumps(output))
