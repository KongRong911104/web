# pip3 install pymysql fake-useragent urllib3 beautifulsoup4 openpyxl python-docx pandas
import os
import sys
import json as js
from fake_useragent import UserAgent
import pandas as pd
import docx
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import time as t
import datetime
import pymysql
t0=t.time()
def time_error(y,s):
    y=int(y)
    currentDateTime = datetime.datetime.now()
    date = currentDateTime.date()
    now_year = int(date.strftime("%Y"))-1911
    if y>now_year or y<now_year-5:
        return 1
    if s != "上" and s!="下":
        return 2
    return 3
def update_table(location,text,chinese=0,color=0,center=0):
    location.text=text
    if center==1:
        location.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if chinese==1:
        paragraph = location.paragraphs[0]
        run = paragraph.runs[0]
        # 获取运行的_element对象
        run_element = run._element
        rpr = run_element.get_or_add_rPr()
        rFonts = rpr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), "標楷體")
    else:
        location.paragraphs[0].style.font.name = 'Times New Roman'
    for run in location.paragraphs[0].runs:
        if color == 2:
            run.font.color.rgb = RGBColor(255, 0, 0)
        else:
            run.font.color.rgb = RGBColor(0, 0, 255)
        if color == 1:
            run.font.size = Pt(13)
        if color == 4:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)
        else:
            run.font.size = Pt(18)
def is_int(number):
    if number == int(number):
        output = int(number)
    else:
        output = float(number)
    return output
ua = UserAgent()
t0=t.time()
year = str(sys.argv[2])
semester = sys.argv[3]
ter=time_error(year,semester)
if ter==1:
    print(js.dumps("年分錯誤"))
    sys.exit()
if ter==2:
    print(js.dumps("學期錯誤" ))
    sys.exit()

if semester == "上":
    sem = "1"
elif semester == "下":
    sem = "2"

cs = {}
input_file="./spider/"+str(year)+"_"+sem+".txt"
with open(input_file, "r",encoding="utf-8") as file:
    for line in file:
        key, values = line.strip().split(":")
        value_list = values.split(",")
        cs[int(key)] = value_list


file = sys.argv[1].split("::123::")[:-1]
for i in range(len(file)):
    if file[i].split('.')[-1] != 'xlsx':
        print(js.dumps("副檔名格式錯誤 {}".format(file[i].split('.')[-1])))
        sys.exit()

readExcelFiles = []
for f in file:
    df = pd.read_excel(f, usecols=[7, 11, 10, 15, 17], engine='openpyxl', dtype={
                       "score": 'int32', "curr_code": 'int32', "cour_cname": 'str', "curr_cname": 'str', 'teac_name': 'str'})
    readExcelFiles.append(df)
    os.remove(f)
readExcelFile = pd.concat(readExcelFiles, axis=0)

data = []
for curr_code in readExcelFile["curr_code"].unique():
    data.append(readExcelFile.loc[readExcelFile['curr_code'] == curr_code])
del readExcelFiles
a = []


for i in range(len(data)):
    curr_code = data[i]['curr_code'].iloc[0]
    cour_cname = data[i]['cour_cname'].iloc[0]
    curr_cname = data[i]['curr_cname'].iloc[0]
    teac_name = data[i]['teac_name'].iloc[0].split(' ')[-1][0:3]
    avg_score = data[i].loc[data[i]['score'] > 20]
    avg_score = round(avg_score['score'].mean(), 1)
    num = len(data)
    under_60 = data[i].loc[data[i]['score'] < 60]
    under_20 = under_60.loc[under_60['score'] <= 20]
    a.append([curr_code,cour_cname,curr_cname,teac_name,avg_score,num,len(under_60),len(under_20)])
del data
a=sorted(a, key=lambda x: (x[3], x[0]))
b=set([i[3] for i in a])
w=[]
for j in b:
    m=[]
    # print(j)
    for k in range(len(a)):
        if a[k][3]==j:
            m.append(a[k])
    w.append(m)
w=sorted(w,key=lambda x: x[0][0])
a = [item for sublist in w for item in sublist]
del w
del b
def format_number(num):
    if num.is_integer():
        return int(num)
    else:
        return num
import math as m
shi_index=1
init_teac_name = a[0][3]
fin_docx=docx.Document()
fin_docx2=docx.Document()
sn=1
sn_node={}


def connection():
    server = "localhost"           # 伺服器
    dbuser = "kongrong"            # 使用者帳號
    dbpassword = "kongrong1104"    # 使用者密碼
    dbname = "test2"                 # 資料庫名稱
    return pymysql.connect(host=server, user=dbuser, password=dbpassword, database=dbname)


def rough_read(from_table):
    conn = connection()  # 建立資料庫連線
    with conn.cursor() as cursor:
        # 查詢資料SQL語法
        command = "SELECT * FROM {}".format(from_table)
        # 執行指令
        cursor.execute(command)
        # 取得所有資料
        result = cursor.fetchall()
        return result
    # return 


exdata=rough_read("exdata")
# print(classname[0])
class_name={}
for i in range(len(exdata)):
    class_name[exdata[i][1]]=exdata[i][2:]

for i in range(len(a)):
    curr_code = a[i][0]
    compulsory = cs[curr_code][0]
    academic = cs[curr_code][1]
    cour_cname = a[i][1]
    curr_cname = a[i][2]
    teac_name = a[i][3]
    avg_score = a[i][4]
    num = a[i][5]
    under_60 = a[i][6]
    under_20 = a[i][7]
    if compulsory[0] == "必" :
        if cour_cname not in sn_node:
            sn_node[cour_cname]="{}-1".format(sn)
            sn+=1
        else:
            sn_node[cour_cname]=sn_node[cour_cname].split("-")[0]+"-"+str(int(sn_node[cour_cname].split("-")[1])+1)
        if teac_name!= init_teac_name:
            shi_index=1
            init_teac_name=teac_name
        math = class_name[cour_cname][1]
        base = class_name[cour_cname][2]
        li = class_name[cour_cname][3]
        shi = class_name[cour_cname][4]
        total_mbls=math+base+li+shi
        math=is_int(math)
        base=is_int(base)
        li=is_int(li)
        shi=is_int(shi)
        if math==0:
            math=""
        if base==0:
            base=""
        if li==0:
            li=""
        if shi==0:
            shi=""
        mh=class_name[cour_cname][0]
        value1 = class_name[cour_cname][5]
        value2 = class_name[cour_cname][6]
        value3 = class_name[cour_cname][7]
        value4 = class_name[cour_cname][8]
        value5 = class_name[cour_cname][9]
        value6 = class_name[cour_cname][10]
        value7 = class_name[cour_cname][11]
        value8 = class_name[cour_cname][12]
        if value1=="NULL":
            value1=""
        if value2=="NULL":
            value2=""
        if value3=="NULL":
            value3=""
        if value4=="NULL":
            value4=""
        if value5=="NULL":
            value5=""
        if value6=="NULL":
            value6=""
        if value7=="NULL":
            value7=""
        if value8=="NULL":
            value8=""
        doc2=docx.Document()
        if academic=="日間部":
            doc2=docx.Document("./反思大學.docx")
        else:
            doc2=docx.Document("./反思碩士.docx")
        paragraphs=doc2.paragraphs
        paragraphs[0].text="{}學年度整合性專題實作課程(Capstone)之課程分析及反思表".format(str(year))
        tables2=doc2.tables[0]
        tables2.cell(3,0).text=sn_node[cour_cname]
        update_table(tables2.cell(3,1),cour_cname,1,4,1)
        update_table(tables2.cell(3,2),academic,1,4,1)
        update_table(tables2.cell(3,3),teac_name,1,4,1)
        update_table(tables2.cell(3,4),curr_cname,1,4,1)
        update_table(tables2.cell(3,5),str(int(total_mbls)),0,4,1)
        update_table(tables2.cell(3,6),str(math),0,4,1)
        update_table(tables2.cell(3,7),str(base),0,4,1)
        update_table(tables2.cell(3,8),str(li),0,4,1)
        update_table(tables2.cell(3,9),str(shi),0,4,1)
        update_table(tables2.cell(3,10),str(mh),0,4,1)
        tables2.cell(3,11).text=value1
        tables2.cell(3,12).text=value2
        tables2.cell(3,13).text=value3
        tables2.cell(3,14).text=value4
        tables2.cell(3,15).text=value5
        tables2.cell(3,16).text=value6
        tables2.cell(3,17).text=value7
        if academic=="日間部":
            tables2.cell(3,18).text=value8
            tables2.cell(3,19).text=str(num)
            update_table(tables2.cell(3,21),str(avg_score),0,4,1)
            update_table(tables2.cell(3,22),str(100-(round((under_60/num)*100,1))),0,4,1)
            tables2.cell(4,22).text="(教師對課程之反思) \n 1."
            paragraph = tables2.cell(4,22).paragraphs[0]
        else:
            tables2.cell(3,18).text=str(num)
            update_table(tables2.cell(3,20),str(avg_score),0,4,1)
            update_table(tables2.cell(3,21),str(100-(round((under_60/num)*100,1))),0,4,1)
            tables2.cell(4,21).text="(教師對課程之反思) \n 1."
            paragraph = tables2.cell(4,21).paragraphs[0]
        shi_index+=1
        run = paragraph.runs[0]
        run.bold=True
        if i < len(a)-1:
            doc2.add_page_break()
        if i == 0:
            fin_docx2=doc2
        else:
            for element in doc2.element.body:
                fin_docx2.element.body.append(element)
    doc = docx.Document('./教學成效表_空.docx')
    tables = doc.tables
    # 老師
    update_table(tables[0].cell(0, 1),teac_name,1,1)
    
    # 學年
    update_table(tables[0].cell(1, 1),year,center=1)

    # 學期
    update_table(tables[0].cell(1, 4),semester,chinese=1,center=1)

    # 課程代號
    update_table(tables[0].cell(1, 6),str(curr_code),center=1)

    # 課程名稱
    update_table(tables[0].cell(2, 1),cour_cname,1,1)

    # 學制
    update_table(tables[0].cell(2, 6),academic,1,1,1)

    # 修習年級
    update_table(tables[0].cell(3, 1),curr_cname,1,1,1)

    # 必/選修
    update_table(tables[0].cell(3, 4),compulsory,1,1,1)

    # 平均
    update_table(tables[0].cell(3, 6),str(avg_score),center=1)

    # 人數
    update_table(tables[0].cell(4, 1),str(num),1,1,1)

    # 60以下
    update_table(tables[0].cell(4, 4),str(under_60),center=1)

    # 20以下
    update_table(tables[0].cell(4, 6),str(under_20),center=1)

    tables[0].cell(5, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    if (under_60/num >= 0.4):
        update_table(tables[0].cell(5, 1),"✓",color=2,center=1)

    if avg_score < 60:
        update_table(tables[0].cell(6, 1),"✓",color=2,center=1)

    if i < len(a)-1:
        doc.add_page_break()
    if i == 0:
        fin_docx=doc
    else:
        for element in doc.element.body:
            fin_docx.element.body.append(element)

# 获取当前日期和时间
now = datetime.datetime.now()

month = now.month
day = now.day
hour = now.hour
minute = now.minute
second = now.second
dday="_{}-{}-{}-{}:{}:{}".format(year,month,day,hour,minute,second)
fin_docx.save("./word_output/{}/{}教學成效表{}.docx".format(sys.argv[4],year+semester,dday))
last_paragraph = fin_docx2.paragraphs[-1]

# 刪除最後一個段落
p = last_paragraph._element
p.getparent().remove(p)
p._p = p._element = None
fin_docx2.save("./word_output/{}/{}教學反思表{}.docx".format(sys.argv[4],year+semester,dday))
with open('xxx.txt', 'w') as f:
    f.write('{}\n'.format(str(t.time()-t0)))
